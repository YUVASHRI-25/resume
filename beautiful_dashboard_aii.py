"""
Beautiful Modern Resume Insight Dashboard (Soft Pastel Rainbow Edition)
-- WITH:
 - Numbered AI Assistant responses (1., 2., 3., ...)
 - Clean pastel AI chat cards
 - Topic-colored AI cards (Skills, ATS, Experience, Keywords)
 - Properly aligned, rounded Suggested Keyword chips
 - Matched Pro Tip pastel box
"""


import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
from typing import Dict, Any, Optional
import io
import base64
import re
import hashlib
import psycopg2
from psycopg2.extras import RealDictCursor
import PyPDF2
import pdfplumber
from docx import Document
import os
import json
from dotenv import load_dotenv
import random
from openai import OpenAI
import requests
import json
import spacy
from spacy.matcher import PhraseMatcher
from nltk.sentiment import vader
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# ...existing code...
# Load environment variables from .env file
load_dotenv('.env', override=True)
# Also try loading from 'env' file if it exists
load_dotenv('env', override=True)

# Insert RAG + helper imports / class here using LangChain
import os
import numpy as np

# LangChain imports
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
    from langchain_community.vectorstores import FAISS
    from langchain_classic.chains import RetrievalQA
    from langchain_core.prompts import PromptTemplate
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ LangChain import error: {e}")
    LANGCHAIN_AVAILABLE = False

class LangChainRAGClient:
    """
    LangChain-based RAG system:
    - LangChain controls the pipeline
    - FAISS for vector storage and retrieval
    - OpenAI for embeddings and generation
    """
    def __init__(self, openai_api_key=None, emb_model="text-embedding-3-small", llm_model="gpt-4o-mini"):
        if not LANGCHAIN_AVAILABLE:
            raise RuntimeError("LangChain not available. Install: pip install langchain langchain-openai langchain-community")
        
        self.api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise RuntimeError("OpenAI API key not found")
        
        self.emb_model = emb_model
        self.llm_model = llm_model
        self.vectorstore = None
        self.qa_chain = None
        
        # Initialize embeddings
        try:
            self.embeddings = OpenAIEmbeddings(
                model=emb_model,
                openai_api_key=self.api_key
            )
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "quota" in error_msg.lower():
                raise RuntimeError("OpenAI quota exceeded. Please add credits to your account at https://platform.openai.com/settings/organization/billing")
            elif "401" in error_msg:
                raise RuntimeError("Invalid OpenAI API key. Please check your env file.")
            else:
                raise RuntimeError(f"OpenAI API error: {error_msg}")
        
        # Initialize LLM
        self.llm = ChatOpenAI(
            model=llm_model,
            temperature=0.3,
            max_tokens=1000,
            openai_api_key=self.api_key
        )
        
        print(f"✅ LangChain RAG initialized with {llm_model}")

    def build_index(self, text: str):
        """
        Build FAISS index from resume text using LangChain
        Args:
            text: Full resume text
        """
        try:
            # Use LangChain's RecursiveCharacterTextSplitter for better chunking
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,  # Characters per chunk
                chunk_overlap=100,  # Overlap between chunks for context
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]  # Try these separators in order
            )
            
            # Split text into chunks
            chunks = text_splitter.split_text(text)
            
            if not chunks:
                raise ValueError("No text chunks created from resume")
            
            # Create Document objects for LangChain
            documents = [Document(page_content=chunk) for chunk in chunks]
            
            # Create FAISS vector store with LangChain
            self.vectorstore = FAISS.from_documents(
                documents=documents,
                embedding=self.embeddings
            )
            
            # Create custom prompt template for resume analysis
            prompt_template = """You are an expert resume analyst. Use the following resume content to answer the question accurately and specifically.

Resume Content:
{context}

Question: {question}

Instructions:
- Answer ONLY based on the information in the resume content provided above
- Be specific and cite exact details from the resume
- If the information is not in the resume, clearly state "This information is not mentioned in the resume"
- Format your response clearly with proper numbering if listing multiple items
- For projects, include: name, description, and technologies
- For certifications, include: name and year (if mentioned)
- For work experience, include: position, company, duration, and key responsibilities

Answer:"""

            PROMPT = PromptTemplate(
                template=prompt_template,
                input_variables=["context", "question"]
            )
            
            # Create RetrievalQA chain with LangChain
            self.qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",  # Stuff all retrieved docs into prompt
                retriever=self.vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}  # Retrieve top 5 most relevant chunks
                ),
                return_source_documents=True,
                chain_type_kwargs={"prompt": PROMPT}
            )
            
            print(f"✅ LangChain RAG index built with {len(chunks)} chunks")
            return True
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ Error building LangChain index: {e}")
            if "429" in error_msg or "quota" in error_msg.lower():
                raise RuntimeError("OpenAI quota exceeded. Please add credits to your account at https://platform.openai.com/settings/organization/billing")
            elif "401" in error_msg:
                raise RuntimeError("Invalid OpenAI API key. Please check your env file.")
            else:
                raise RuntimeError(f"Error building index: {error_msg}")

    def query(self, user_query: str, k=5):
        """
        Query the resume using LangChain RAG pipeline
        Args:
            user_query: Question to ask about the resume
            k: Number of chunks to retrieve (default: 5)
        Returns:
            Answer from the LLM based on retrieved context
        """
        if self.qa_chain is None:
            return "❌ RAG system not initialized. Please upload a resume first."
        
        try:
            # Query using LangChain's RetrievalQA chain
            result = self.qa_chain.invoke({"query": user_query})
            
            answer = result.get("result", "").strip()
            
            # Add source information for transparency
            if result.get("source_documents"):
                num_sources = len(result["source_documents"])
                answer += f"\n\n_[Based on {num_sources} relevant sections from the resume]_"
            
            return answer
            
        except Exception as e:
            error_message = str(e)
            print(f"❌ LangChain query error: {e}")
            
            # Handle specific OpenAI errors with helpful messages
            if "429" in error_message or "quota" in error_message.lower():
                return """🚫 OpenAI API Quota Exceeded

Your OpenAI API key has run out of credits. Here's how to fix this:

1. **Check Your Usage**: Visit https://platform.openai.com/usage
2. **Add Credits**: Go to https://platform.openai.com/settings/organization/billing
3. **Set Up Billing**: Add a payment method and purchase credits ($10 recommended)

The RAG system requires credits to:
- Create embeddings (very cheap)
- Generate answers (affordable with GPT-4o-mini)

Total cost per resume: ~$0.02-0.05"""
            
            elif "401" in error_message or "authentication" in error_message.lower():
                return """🔑 Authentication Error

Your OpenAI API key is invalid or expired. Please:
1. Get a valid key from https://platform.openai.com/api-keys
2. Update your 'env' file: OPENAI_API_KEY=sk-proj-...
3. Restart the application"""
            
            elif "rate_limit" in error_message.lower():
                return """⏱️ Rate Limit Exceeded

You're making too many requests. Please:
1. Wait 60 seconds and try again
2. Consider upgrading your OpenAI plan for higher limits"""
            
            else:
                return f"""❌ Error generating response

Error details: {error_message}

Please try:
1. Rephrasing your question
2. Checking your internet connection
3. Verifying your OpenAI API key has credits"""
    
    def retrieve_context(self, query: str, k=5):
        """Retrieve relevant chunks without generating answer (for debugging)"""
        if self.vectorstore is None:
            return []
        
        try:
            docs = self.vectorstore.similarity_search(query, k=k)
            return [doc.page_content for doc in docs]
        except Exception as e:
            print(f"❌ Retrieval error: {e}")
            return []
    # ...existing code...
# ...existing code...
class AIAssistant:
    """Light wrapper for OpenAI Chat Completions API (graceful if no API key)."""
    def __init__(self, api_key: Optional[str] = None, model: str = "gpt-4o-mini"):
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        self.model = model
        self.enabled = False
        self.client = None
        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key)
                self.enabled = True
                print(f"✅ OpenAI client initialized successfully with model: {self.model}")
            except Exception as e:
                print(f"❌ OpenAI client initialization failed: {e}")
                self.client = None
                self.enabled = False
        else:
            print("⚠️ No OpenAI API key found. Set OPENAI_API_KEY environment variable.")
            self.enabled = False

    def get_response(self, prompt: str, context: Optional[str] = None, max_tokens: int = 512) -> str:
        """Return text or a clear disabled message when API key is missing."""
        if not self.enabled or self.client is None:
            return (
                "🤖 AI Assistant is currently unavailable. "
                "To enable AI features, please set your OPENAI_API_KEY in the environment variables or .env file. "
                "You can get an API key from https://platform.openai.com/api-keys"
            )
        
        try:
            # Prepare the full prompt with context
            full_prompt = f"Context:\n{context}\n\nUser: {prompt}" if context else prompt
            
            # Use the correct OpenAI Chat Completions API
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert resume consultant and career advisor. Provide specific, actionable advice for resume improvement. Format your response as clean, numbered points without any markdown formatting (no **bold**, no # headers, no bullet points). Use simple numbered lists with clear, professional language. Keep responses concise and professional."
                    },
                    {
                        "role": "user", 
                        "content": full_prompt
                    }
                ],
                max_tokens=max_tokens,
                temperature=0.7
            )
            
            # Extract the response content and clean it
            if response and response.choices:
                raw_response = response.choices[0].message.content.strip()
                # Clean the response to remove markdown formatting
                cleaned_response = self._clean_response(raw_response)
                return cleaned_response
            else:
                return "Sorry, I couldn't generate a response. Please try again."
                
        except Exception as e:
            print(f"❌ OpenAI API error: {e}")
            return f"🤖 AI Assistant temporarily unavailable. Error: {str(e)}"
    
    def _clean_response(self, response: str) -> str:
        """Clean AI response to remove markdown formatting and ensure proper numbering."""
        import re
        
        # Remove markdown bold formatting
        response = re.sub(r'\*\*(.*?)\*\*', r'\1', response)
        
        # Remove markdown italic formatting
        response = re.sub(r'\*(.*?)\*', r'\1', response)
        
        # Remove markdown headers
        response = re.sub(r'^#+\s*', '', response, flags=re.MULTILINE)
        
        # Remove bullet points and replace with proper numbering
        response = re.sub(r'^[\s]*[-•]\s*', '', response, flags=re.MULTILINE)
        
        # Clean up any remaining markdown
        response = re.sub(r'`(.*?)`', r'\1', response)
        
        # Ensure proper line breaks
        response = response.replace('\n\n', '\n')
        
        # Clean up extra spaces
        response = re.sub(r'\n\s+', '\n', response)
        
        return response.strip()
# ...existing code...
 
   

# Configure Streamlit page
st.set_page_config(
    page_title="RESUME INSIGHT 🌈",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ---------------------------
# Clean White and Blue CSS Theme
# ---------------------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');


* { font-family: 'Inter', sans-serif; }


/* App background */
.stApp {
    background: linear-gradient(135deg, #ffffff 0%, #f8fafc 50%, #e2e8f0 100%);
    min-height: 100vh;
}


/* Header */
.modern-header {
    background: linear-gradient(120deg, #ffffff 0%, #f1f5f9 50%, #e2e8f0 100%);
    padding: 3rem 2rem;
    border-radius: 20px;
    margin-bottom: 2rem;
    box-shadow: 0 15px 40px rgba(59, 130, 246, 0.1);
    border: 1px solid #e2e8f0;
    position: relative;
}
.header-title { font-size: 3rem; font-weight: 800; color: #1e293b; text-shadow: 0 2px 8px rgba(255,255,255,0.8); }
.header-subtitle { font-size: 1.2rem; color: #475569; font-weight: 600; }


/* Metric Cards */
.metric-card {
    background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem 0;
    text-align: center;
    box-shadow: 0 10px 25px rgba(59, 130, 246, 0.08);
    border: 1px solid #e2e8f0;
    transition: all 0.4s ease;
}
.metric-card:hover { transform: translateY(-5px); box-shadow: 0 20px 40px rgba(59, 130, 246, 0.15); }
.metric-value { font-size: 3rem; font-weight: 800; background: linear-gradient(90deg, #3b82f6 0%, #1d4ed8 50%, #1e40af 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.metric-label { font-size: 1rem; font-weight: 600; color: #64748b; }


/* Alerts */
.alert-modern { border-radius: 12px; padding: 1.5rem 2rem; margin: 1rem 0; font-weight: 500; box-shadow: 0 4px 15px rgba(59, 130, 246, 0.08); }
.alert-success { background: linear-gradient(135deg, #dcfce7 0%, #bbf7d0 100%); color: #14532d; border-left: 4px solid #16a34a; }
.alert-warning { background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%); color: #78350f; border-left: 4px solid #f59e0b; }
.alert-error { background: linear-gradient(135deg, #fee2e2 0%, #fecaca 100%); color: #7f1d1d; border-left: 4px solid #ef4444; }


/* Cards */
.modern-card { background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%); border-radius: 16px; padding: 2rem; margin: 1rem 0; box-shadow: 0 8px 25px rgba(59, 130, 246, 0.08); border: 1px solid #e2e8f0; }


/* Buttons */
.stButton > button {
    background: linear-gradient(120deg, #3b82f6 0%, #1d4ed8 30%, #1e40af 100%);
    color: #ffffff; font-weight: 700; border-radius: 12px; padding: 0.75rem 2rem; border: none; transition: all 0.2s ease;
    box-shadow: 0 8px 0 #1e3a8a, 0 12px 24px rgba(30,58,138,0.35);
    transform: translateY(0);
}
.stButton > button:hover { transform: translateY(-2px); box-shadow: 0 10px 0 #1e3a8a, 0 16px 28px rgba(30,58,138,0.45); }
.stButton > button:active { transform: translateY(2px); box-shadow: 0 6px 0 #1e3a8a, 0 8px 16px rgba(30,58,138,0.4); }


/* Tabs */
.stTabs [data-baseweb="tab-list"] { background: linear-gradient(90deg, #f1f5f9 0%, #e2e8f0 100%); border-radius: 16px; padding: 0.5rem; box-shadow: 0 4px 15px rgba(59, 130, 246, 0.08); }
.stTabs [aria-selected="true"] { background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); color: #ffffff !important; border-radius: 12px; box-shadow: 0 6px 20px rgba(59, 130, 246, 0.2); }


/* Sidebar */
.sidebar-modern { background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%); border-radius: 16px; padding: 1.5rem; box-shadow: 0 6px 20px rgba(59, 130, 246, 0.08); border: 1px solid #e2e8f0; }


/* Skill Tags */
.skill-tag {
    display: inline-block; background: linear-gradient(135deg, #dbeafe 0%, #bfdbfe 100%); color: #1e40af; border-radius: 20px; padding: 0.5rem 1rem; margin: 0.3rem; font-weight: 600; box-shadow: 0 4px 12px rgba(59, 130, 246, 0.1); transition: all 0.3s ease; border: 1px solid #93c5fd;
}
.skill-tag:hover { transform: translateY(-2px); background: linear-gradient(135deg, #bfdbfe 0%, #93c5fd 100%); }


/* Suggested Keywords (chips) */
.keyword-suggested {
    display: inline-block; background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); color: #0c4a6e; border-radius: 25px; padding: 0.6rem 1.2rem; margin: 0.4rem; font-weight: 600; font-size: 0.9rem; white-space: nowrap; box-shadow: 0 4px 10px rgba(59, 130, 246, 0.1); transition: transform 0.3s ease, background 0.3s ease; border: 1px solid #bae6fd;
}
.keyword-suggested:hover { transform: translateY(-3px); background: #bae6fd; }


/* Progress bar */
.stProgress > div > div > div { background: linear-gradient(90deg, #3b82f6 0%, #1d4ed8 25%, #1e40af 50%, #1e3a8a 75%, #172554 100%); border-radius: 10px; }


/* Chat messages (base) */
.chat-message { border-radius: 16px; padding: 1.5rem; margin: 0.75rem 0; font-size: 0.95rem; line-height: 1.8; box-shadow: 0 6px 20px rgba(59, 130, 246, 0.08); }


/* Chat - default AI */
.chat-ai { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 50%, #dbeafe 100%); color: #1e293b; border-left: 6px solid #3b82f6; }


/* Chat - user */
.chat-user { background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 50%, #e2e8f0 100%); color: #1e293b; border-left: 6px solid #64748b; text-align:left; }


/* Topic-specific AI colors */
.chat-ai-skills { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 50%, #dbeafe 100%); border-left: 6px solid #3b82f6; }
.chat-ai-ats { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 50%, #dbeafe 100%); border-left: 6px solid #0ea5a0; }
.chat-ai-experience { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 50%, #dbeafe 100%); border-left: 6px solid #16a34a; }
.chat-ai-keywords { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 50%, #dbeafe 100%); border-left: 6px solid #7c3aed; }


/* Chat headings/emphasis */
.chat-message strong { font-weight: 700; color: #0f172a; display: block; margin-bottom: 0.5rem; }
.chat-ai div { line-height: 1.9; margin-left: 0.5rem; }


/* Section header style */
.section-header-modern { background: linear-gradient(120deg, #3b82f6 0%, #1d4ed8 30%, #1e40af 100%); color: #ffffff; font-weight: 700; text-align: center; padding: 1rem; border-radius: 16px; margin-bottom: 1.5rem; box-shadow: 0 8px 20px rgba(59, 130, 246, 0.2); }


/* Pro Tip box (info) */
.alert-modern.alert-info { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); color: #0c4a6e; border-left: 6px solid #3b82f6; box-shadow: 0 6px 20px rgba(59, 130, 246, 0.1); padding: 1rem 1.5rem; border-radius: 12px; font-size: 0.95rem; line-height: 1.7; }


/* Login container */
.login-container { background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%); border-radius: 20px; padding: 3rem; margin: 2rem auto; box-shadow: 0 15px 40px rgba(59, 130, 246, 0.1); border: 1px solid #e2e8f0; }


/* Scrollbar */
::-webkit-scrollbar { width: 8px; }
::-webkit-scrollbar-thumb { background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); border-radius: 10px; }
::-webkit-scrollbar-thumb:hover { background: linear-gradient(135deg, #1d4ed8 0%, #1e40af 100%); }
</style>
""", unsafe_allow_html=True)


# ---------------------------
# Data / DB / Keywords (unchanged logic)
# ---------------------------
DB_CONFIG = {
    'host': 'localhost',
    'database': 'resume_insight',
    'user': 'postgres',
    'password': 'yuva',
    'port': '5432'
}


JOB_KEYWORDS = {
    "Data Scientist": [
        "python", "machine learning", "statistics", "pandas", "numpy", "scikit-learn",
        "tensorflow", "pytorch", "sql", "data analysis", "visualization", "jupyter",
        "r", "statistics", "deep learning", "neural networks", "regression",
        "classification", "clustering", "feature engineering", "model deployment"
    ],
    "Software Engineer": [
        "programming", "java", "python", "javascript", "react", "node.js", "database",
        "git", "agile", "testing", "debugging", "api", "frontend", "backend",
        "algorithms", "data structures", "microservices", "docker", "kubernetes",
        "rest api", "graphql", "version control", "ci/cd"
    ],
    "Product Manager": [
        "product", "strategy", "roadmap", "stakeholder", "analytics", "user experience",
        "market research", "agile", "scrum", "requirements", "metrics", "wireframes",
        "user stories", "product development", "customer research", "competitive analysis",
        "go-to-market", "product launch"
    ],
    "Marketing Manager": [
        "marketing", "digital marketing", "seo", "social media", "analytics", "campaigns",
        "brand", "content", "advertising", "growth", "conversion", "roi", "crm",
        "email marketing", "ppc", "content marketing", "brand management", "lead generation"
    ],
    "Data Analyst": [
        "sql", "excel", "python", "tableau", "power bi", "statistics", "reporting",
        "data visualization", "business intelligence", "analytics", "dashboards", "kpi",
        "data mining", "etl", "data modeling", "statistical analysis", "forecasting"
    ],
    "DevOps Engineer": [
        "docker", "kubernetes", "aws", "azure", "gcp", "jenkins", "ci/cd", "terraform",
        "ansible", "monitoring", "linux", "bash", "infrastructure", "deployment",
        "cloud computing", "automation", "configuration management", "security"
    ],
    "UI/UX Designer": [
        "figma", "sketch", "adobe xd", "prototyping", "wireframes", "user research",
        "usability testing", "design systems", "responsive design", "accessibility",
        "typography", "visual design", "interaction design", "user interface", "user experience"
    ],
    "Cybersecurity Analyst": [
        "security", "penetration testing", "vulnerability assessment", "siem", "firewall",
        "incident response", "compliance", "risk assessment", "cryptography", "network security",
        "security monitoring", "threat analysis", "security policies", "audit"
    ],
    "Business Analyst": [
        "requirements gathering", "process improvement", "stakeholder management",
        "documentation", "business process", "gap analysis", "user stories", "workflow",
        "project management", "data analysis", "business intelligence", "process mapping"
    ],
    "Full Stack Developer": [
        "html", "css", "javascript", "react", "angular", "vue", "node.js", "express",
        "mongodb", "postgresql", "rest api", "graphql", "version control", "responsive design",
        "frontend", "backend", "database", "api development", "web development"
    ],
    "Machine Learning Engineer": [
        "python", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "machine learning",
        "deep learning", "neural networks", "computer vision", "nlp", "data science",
        "algorithms", "statistics", "linear algebra", "calculus", "regression", "classification",
        "clustering", "feature engineering", "model deployment", "mlops", "docker", "kubernetes"
    ],
    "AI Engineer": [
        "artificial intelligence", "machine learning", "deep learning", "neural networks",
        "python", "tensorflow", "pytorch", "computer vision", "nlp", "natural language processing",
        "opencv", "transformers", "bert", "gpt", "reinforcement learning", "generative ai",
        "llm", "chatbot", "model optimization", "ai ethics", "edge ai", "quantization"
    ]
}


# Enhanced Technical Skills Database
TECHNICAL_SKILLS = {
    "programming_languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", "rust",
        "swift", "kotlin", "scala", "matlab", "perl", "shell", "bash", "powershell"
    ],
    "web_technologies": [
        "html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask",
        "spring", "laravel", "rails", "next.js", "nuxt.js", "svelte", "jquery", "bootstrap",
        "tailwind", "sass", "less", "webpack", "vite", "npm", "yarn"
    ],
    "databases": [
        "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra",
        "dynamodb", "oracle", "sqlite", "mariadb", "neo4j", "influxdb", "couchdb"
    ],
    "cloud_platforms": [
        "aws", "azure", "gcp", "heroku", "digital ocean", "linode", "vultr", "cloudflare",
        "firebase", "vercel", "netlify", "railway"
    ],
    "devops_tools": [
        "docker", "kubernetes", "jenkins", "gitlab ci", "github actions", "terraform",
        "ansible", "chef", "puppet", "vagrant", "prometheus", "grafana", "elk stack"
    ],
    "data_science": [
        "machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn",
        "pandas", "numpy", "matplotlib", "seaborn", "plotly", "jupyter", "rstudio",
        "spark", "hadoop", "kafka", "airflow", "dbt", "mlflow"
    ],
    "mobile_development": [
        "react native", "flutter", "xamarin", "ionic", "cordova", "android studio",
        "xcode", "swift", "kotlin", "objective-c"
    ],
    "testing_tools": [
        "selenium", "cypress", "jest", "mocha", "chai", "pytest", "junit", "testng",
        "cucumber", "postman", "insomnia", "k6", "jmeter"
    ],
    "design_tools": [
        "figma", "sketch", "adobe xd", "photoshop", "illustrator", "canva", "invision",
        "zeplin", "principle", "framer"
    ],
    "project_management": [
        "jira", "confluence", "trello", "asana", "monday.com", "notion", "slack",
        "microsoft teams", "zoom", "github", "gitlab", "bitbucket"
    ],
    "business_intelligence": [
        "tableau", "power bi", "looker", "qlikview", "excel", "google sheets",
        "salesforce", "hubspot", "zapier"
    ]
}


# Enhanced Soft Skills Database
SOFT_SKILLS = {
    "leadership": [
        "leadership", "team leadership", "team management", "mentoring", "coaching",
        "people management", "strategic leadership", "cross-functional leadership"
    ],
    "communication": [
        "communication", "verbal communication", "written communication", "presentation",
        "public speaking", "technical writing", "documentation", "client communication"
    ],
    "collaboration": [
        "teamwork", "collaboration", "cross-functional collaboration", "stakeholder management",
        "interpersonal skills", "relationship building", "networking"
    ],
    "problem_solving": [
        "problem solving", "critical thinking", "analytical thinking", "troubleshooting",
        "debugging", "root cause analysis", "creative problem solving"
    ],
    "project_management": [
        "project management", "agile", "scrum", "kanban", "time management", "deadline management",
        "resource management", "risk management", "change management"
    ],
    "adaptability": [
        "adaptability", "flexibility", "resilience", "learning agility", "continuous learning",
        "innovation", "creativity", "open-mindedness"
    ],
    "customer_focus": [
        "customer service", "customer satisfaction", "user experience", "client relations",
        "customer success", "support", "helpdesk"
    ],
    "business_acumen": [
        "business acumen", "strategic thinking", "commercial awareness", "market analysis",
        "competitive analysis", "business development", "sales", "negotiation"
    ]
}


# ---------------------------
# Database Manager (unchanged)
# ---------------------------
class DatabaseManager:
    def __init__(self):
        self.config = DB_CONFIG


    def get_connection(self):
        try:
            return psycopg2.connect(**self.config)
        except Exception as e:
            st.error(f"Database connection failed: {e}")
            return None


    def create_tables(self):
        conn = self.get_connection()
        if not conn:
            return False
        try:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS users (
                        id SERIAL PRIMARY KEY,
                        username VARCHAR(50) UNIQUE NOT NULL,
                        email VARCHAR(100) UNIQUE NOT NULL,
                        password_hash VARCHAR(255) NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS resume_analyses (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER REFERENCES users(id),
                        filename VARCHAR(255) NOT NULL,
                        job_role VARCHAR(100) NOT NULL,
                        file_hash VARCHAR(64) NOT NULL,
                        ats_score DECIMAL(5,2),
                        role_match_percentage DECIMAL(5,2),
                        overall_score DECIMAL(5,2),
                        technical_skills TEXT[],
                        soft_skills TEXT[],
                        found_keywords TEXT[],
                        missing_keywords TEXT[],
                        analysis_data TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(user_id, file_hash, job_role)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS chat_history (
                        id SERIAL PRIMARY KEY,
                        user_id INTEGER REFERENCES users(id),
                        session_id VARCHAR(100),
                        message TEXT NOT NULL,
                        response TEXT NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                conn.commit()
                return True
        except Exception as e:
            st.error(f"Error creating tables: {e}")
            return False
        finally:
            conn.close()


    def authenticate_user(self, username, password):
        conn = self.get_connection()
        if not conn:
            return None
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                password_hash = hashlib.sha256(password.encode()).hexdigest()
                cur.execute(
                    "SELECT id, username, email FROM users WHERE username = %s AND password_hash = %s",
                    (username, password_hash)
                )
                user = cur.fetchone()
                return dict(user) if user else None
        except Exception as e:
            st.error(f"Authentication error: {e}")
            return None
        finally:
            conn.close()


    def register_user(self, username, email, password):
        conn = self.get_connection()
        if not conn:
            return False
        try:
            with conn.cursor() as cur:
                password_hash = hashlib.sha256(password.encode()).hexdigest()
                cur.execute(
                    "INSERT INTO users (username, email, password_hash) VALUES (%s, %s, %s)",
                    (username, email, password_hash)
                )
                conn.commit()
                return True
        except Exception as e:
            st.error(f"Registration error: {e}")
            return False
        finally:
            conn.close()


    def save_analysis(self, user_id, filename, job_role, analysis_data, file_content=None):
        conn = self.get_connection()
        if not conn:
            return False
        try:
            with conn.cursor() as cur:
                # Generate file hash to prevent duplicates
                file_hash = hashlib.sha256(file_content.encode() if file_content else filename.encode()).hexdigest()
               
                analysis_json = json.dumps(analysis_data)
                cur.execute("""
                    INSERT INTO resume_analyses 
                    (user_id, filename, job_role, file_hash, ats_score, role_match_percentage, overall_score,
                     technical_skills, soft_skills, found_keywords, missing_keywords, analysis_data)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (user_id, file_hash, job_role)
                    DO UPDATE SET
                        filename = EXCLUDED.filename,
                        ats_score = EXCLUDED.ats_score,
                        role_match_percentage = EXCLUDED.role_match_percentage,
                        overall_score = EXCLUDED.overall_score,
                        technical_skills = EXCLUDED.technical_skills,
                        soft_skills = EXCLUDED.soft_skills,
                        found_keywords = EXCLUDED.found_keywords,
                        missing_keywords = EXCLUDED.missing_keywords,
                        analysis_data = EXCLUDED.analysis_data,
                        created_at = CURRENT_TIMESTAMP
                """, (
                    user_id, filename, job_role, file_hash,
                    analysis_data.get('ats_score'),
                    analysis_data.get('role_match_percentage'),
                    analysis_data.get('overall_score'),
                    analysis_data.get('technical_skills', []),
                    analysis_data.get('soft_skills', []),
                    analysis_data.get('found_keywords', []),
                    analysis_data.get('missing_keywords', []),
                    analysis_json
                ))
                conn.commit()
                return True
        except Exception as e:
            st.error(f"Error saving analysis: {e}")
            return False
        finally:
            conn.close()


    def get_user_analyses(self, user_id):
        conn = self.get_connection()
        if not conn:
            return []
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT * FROM resume_analyses WHERE user_id = %s ORDER BY created_at DESC", (user_id,))
                return [dict(row) for row in cur.fetchall()]
        except Exception as e:
            st.error(f"Error fetching analyses: {e}")
            return []
        finally:
            conn.close()


    def save_chat_message(self, user_id, session_id, message, response):
        conn = self.get_connection()
        if not conn:
            return False
        try:
            with conn.cursor() as cur:
                cur.execute("INSERT INTO chat_history (user_id, session_id, message, response) VALUES (%s, %s, %s, %s)",
                            (user_id, session_id, message, response))
                conn.commit()
                return True
        except Exception as e:
            st.error(f"Error saving chat message: {e}")
            return False
        finally:
            conn.close()


# ---------------------------
# Job Description Analyzer
# ---------------------------
class JobDescriptionAnalyzer:
    def __init__(self):
        self.job_boards_api = {
            'indeed': 'https://api.indeed.com/v1/jobs',
            'linkedin': 'https://api.linkedin.com/v2/jobs',
            'glassdoor': 'https://api.glassdoor.com/api/api.htm'
        }
   
    def extract_job_description_from_text(self, text):
        """Extract job description from pasted text"""
        if not text or len(text.strip()) < 50:
            return None
       
        # Look for common job description patterns
        patterns = [
            r'(?:job description|position|role|responsibilities)[:\s]*(.*?)(?=requirements|qualifications|skills|education|experience)',
            r'(?:about the role|about this role)[:\s]*(.*?)(?=what we offer|benefits|compensation)',
            r'(?:what you\'ll do|key responsibilities)[:\s]*(.*?)(?=what we\'re looking for|requirements)'
        ]
       
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
       
        return text.strip()
   
    def analyze_job_description(self, job_description):
        """Analyze job description to extract key requirements"""
        if not job_description:
            return None
       
        text_lower = job_description.lower()
       
        # Extract required skills
        required_skills = []
        for category, skills in TECHNICAL_SKILLS.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    required_skills.append(skill)
       
        # Extract soft skills
        required_soft_skills = []
        for category, skills in SOFT_SKILLS.items():
            for skill in skills:
                skill_words = skill.lower().split()
                if all(word in text_lower for word in skill_words):
                    required_soft_skills.append(skill)
       
        # Extract experience level
        experience_patterns = {
            'entry': ['entry level', 'junior', '0-2 years', '1-2 years', 'fresh graduate'],
            'mid': ['mid level', '3-5 years', '4-6 years', 'intermediate'],
            'senior': ['senior', 'lead', '5+ years', '6+ years', '7+ years', 'expert'],
            'executive': ['director', 'vp', 'c-level', 'executive', '10+ years']
        }
       
        experience_level = 'mid'  # default
        for level, patterns in experience_patterns.items():
            
            if any(pattern in text_lower for pattern in patterns):
                experience_level = level
                break
       
        # Extract education requirements
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'diploma', 'certification']
        education_required = any(keyword in text_lower for keyword in education_keywords)
       
        return {
            'required_skills': list(set(required_skills)),
            'required_soft_skills': list(set(required_soft_skills)),
            'experience_level': experience_level,
            'education_required': education_required,
            'word_count': len(job_description.split()),
            'raw_text': job_description
        }
   
    def get_job_description_from_url(self, url):
        """Extract job description from URL (placeholder for future API integration)"""
        try:
            # This would integrate with job board APIs
            # For now, return None to indicate manual input needed
            return None
        except Exception as e:
            st.warning(f"Could not fetch job description from URL: {str(e)}")
            return None


# ---------------------------
# Resume Analyzer (unchanged logic)
# ---------------------------
class ResumeAnalyzer:
    def __init__(self):
        self.stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        # Load NLP models once
        try:
            self.nlp = spacy.load('en_core_web_sm')
        except Exception:
            # If model not present, attempt to download quickly
            from spacy.cli import download as spacy_download
            spacy_download('en_core_web_sm')
            self.nlp = spacy.load('en_core_web_sm')
        # Build phrase matchers for skills
        self.tech_matcher = PhraseMatcher(self.nlp.vocab, attr='LOWER')
        self.soft_matcher = PhraseMatcher(self.nlp.vocab, attr='LOWER')
        tech_phrases = []
        for _, skills in TECHNICAL_SKILLS.items():
            tech_phrases.extend(skills)
        soft_phrases = []
        for _, skills in SOFT_SKILLS.items():
            soft_phrases.extend(skills)
        
        self.tech_matcher.add('TECH', [self.nlp.make_doc(s) for s in dict.fromkeys(tech_phrases)])
        self.soft_matcher.add('SOFT', [self.nlp.make_doc(s) for s in dict.fromkeys(soft_phrases)])

        # Initialize AI assistant for on-demand responses
        self.ai_assistant = AIAssistant()
        
        # Initialize LangChain RAG client for resume question answering
        try:
            self.rag_client = LangChainRAGClient(openai_api_key=os.getenv("OPENAI_API_KEY"))
            self.rag_enabled = True
            print("✅ LangChain RAG client initialized successfully")
        except Exception as e:
            print(f"⚠️ LangChain RAG client initialization failed: {e}")
            self.rag_client = None
            self.rag_enabled = False
# ...existing code...

    def extract_text_from_pdf(self, file):
        try:
            with pdfplumber.open(file) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text
        except:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
            except:
                return "Error extracting PDF text"


    def extract_text_from_docx(self, file):
        try:
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except:
            return "Error extracting DOCX text"


    def extract_text_from_txt(self, file):
        try:
            return str(file.read(), "utf-8")
        except:
            return "Error extracting TXT text"

    def build_resume_rag_index(self, text: str):
        """
        Builds a LangChain RAG index from resume text for question answering.
        LangChain handles chunking, embedding, and FAISS indexing automatically.
        
        Args:
            text: The full resume text
        Returns:
            tuple: (success: bool, error_message: str or None)
        """
        if not self.rag_enabled or not self.rag_client:
            print("⚠️ LangChain RAG not enabled, skipping index building")
            return False, "LangChain RAG client not initialized"
        
        try:
            # LangChain handles all the chunking and indexing
            self.rag_client.build_index(text)
            print("✅ LangChain RAG index built successfully")
            return True, None
                
        except RuntimeError as e:
            error_msg = str(e)
            print(f"❌ Error building LangChain RAG index: {e}")
            if "quota exceeded" in error_msg.lower():
                return False, "quota_exceeded"
            elif "Invalid" in error_msg or "401" in error_msg:
                return False, "invalid_key"
            else:
                return False, error_msg
        except Exception as e:
            print(f"❌ Error building LangChain RAG index: {e}")
            return False, str(e)

# ...existing code...
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """
        Robust contact extractor: finds emails, phone numbers, LinkedIn/GitHub/website links,
        and returns a small summary dict with 'found' boolean.
        """
        import re

        if not text:
            return {"emails": [], "phones": [], "linkedin": [], "github": [], "websites": [], "found": False}

        # Consider top lines as likely header with contact info
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        header_block = " ".join(lines[:10])  # top 10 lines

        # Regex patterns
        email_re = re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+')
        phone_re = re.compile(r'(\+?\d{1,3}[\s\-.\(]?\d{1,4}[\s\-.\)]?\d{1,4}[\s\-.\)]?\d{1,9})')
        linkedin_re = re.compile(r'https?://(?:www\.)?linkedin\.com/[\w\-/]+', re.IGNORECASE)
        github_re = re.compile(r'https?://(?:www\.)?github\.com/[\w\-/]+', re.IGNORECASE)
        url_re = re.compile(r'https?://[^\s,;]+', re.IGNORECASE)

        emails = list(dict.fromkeys(email_re.findall(text)))
        phones = list(dict.fromkeys([p.strip() for p in phone_re.findall(text)]))
        linkedin = list(dict.fromkeys(linkedin_re.findall(text)))
        github = list(dict.fromkeys(github_re.findall(text)))
        websites = [u for u in dict.fromkeys(url_re.findall(text)) if not (u in linkedin or u in github)]

        # Heuristic: contact exists if any explicit contact tokens or patterns appear in header block or anywhere
        contact_tokens = re.search(r'\b(contact|email|phone|mobile|tel|linkedin|github)\b', header_block, re.IGNORECASE)
        found = bool(emails or phones or linkedin or github or contact_tokens)

        return {
            "emails": emails,
            "phones": phones,
            "linkedin": linkedin,
            "github": github,
            "websites": websites,
            "found": found
        }


    def extract_sections(self, text):
        """
        Split resume into simple sections by heading heuristics and also populate
        a structured contact_info using extract_contact_info().
        """
        import re
        sections = {}
        section_patterns = {
            'contact': r'(contact|email|phone|mobile|tel|address|linkedin|github)',
            'summary': r'(summary|objective|profile|about|overview)',
            'experience': r'(experience|employment|work|career|professional|job|position)',
            'education': r'(education|academic|qualification|degree|university|college)',
            'skills': r'(skills|technical|competencies|expertise|abilities|technologies)',
            'projects': r'(projects|portfolio|work samples|personal projects)',
            'certifications': r'(certifications?|certificates?|licensed?|credentials)'
        }

        lines = text.splitlines()
        current_section = "header"
        sections[current_section] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            # if a line starts with a section heading word, switch section
            matched = False
            for section_name, pattern in section_patterns.items():
                if re.search(r'^\s*' + pattern + r'\b', stripped, re.IGNORECASE):
                    current_section = section_name
                    sections.setdefault(current_section, [])
                    matched = True
                    break
            if not matched:
                sections.setdefault(current_section, []).append(stripped)

        # Join lists into text blocks
        for k in list(sections.keys()):
            sections[k] = "\n".join(sections[k]).strip()

        # Ensure contact info is populated even if no 'Contact' heading exists
        contact_info = self.extract_contact_info(text)
        if contact_info["found"]:
            # attach structured contact info
            sections["contact_info"] = contact_info
            # if a contact block exists, keep both raw block and structured info
            if not sections.get("contact"):
                sections["contact"] = ""
        else:
            sections.setdefault("contact", "")
            sections["contact_info"] = contact_info

        return sections

# ...existing code...
    # ...existing code...
    def extract_skills_from_sections_only(self, text, sections=None):
        """
        Extract skills ONLY from dedicated 'Skills', 'Technical Skills', or 'Soft Skills' sections.
        Does not extract from project descriptions, experience, or summary sections.
        Returns: (technical_skills, soft_skills) as comma-separated strings
        """
        import re

        if not text:
            return "", ""

        # Get sections if not provided
        if sections is None:
            sections = self.extract_sections(text)

        # Look for skills sections with various naming patterns
        skills_sections = []
        
        # Check for different section names
        section_patterns = {
            'skills': r'(skills|technical skills|soft skills|competencies|expertise|abilities|technologies)',
            'technical_skills': r'(technical skills|technical competencies|programming|languages|tools|frameworks)',
            'soft_skills': r'(soft skills|interpersonal skills|leadership skills|communication skills|management skills)'
        }
        
        # Extract all skills-related sections
        for section_name, pattern in section_patterns.items():
            if section_name in sections and sections[section_name].strip():
                skills_sections.append(sections[section_name].strip())
        
        # If no dedicated skills sections found, return empty
        if not skills_sections:
            return "", ""
        
        # Combine all skills sections
        combined_skills_text = "\n".join(skills_sections)
        
        # Create NLP document for the skills sections only
        doc = self.nlp(combined_skills_text)
        text_lower = combined_skills_text.lower()

        raw_tech = []
        raw_soft = []

        # 1) PhraseMatcher results (high precision)
        for _id, start, end in self.tech_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_tech.append(span)
        for _id, start, end in self.soft_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_soft.append(span)

        # token set for exact matches
        token_texts = [t.text for t in doc if t.text.strip()]
        token_set = {t.lower() for t in token_texts}

        # Flatten master lists
        def flatten_master(master):
            out = []
            if isinstance(master, dict):
                for _, lst in master.items():
                    out.extend(lst)
            elif isinstance(master, list):
                out.extend(master)
            return out

        tech_master = flatten_master(TECHNICAL_SKILLS)
        soft_master = flatten_master(SOFT_SKILLS)

        # helpers
        def already_found(raw_list, skill):
            return skill.strip().lower() in {s.strip().lower() for s in raw_list}

        def add_if_found(raw_list, skill):
            if not already_found(raw_list, skill):
                raw_list.append(skill)

        # 2) Exact token matching for short/special skills (prevent substring noise)
        for skill in tech_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            # skills containing punctuation (C++, C#, node.js) -> regex whole-word check
            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, combined_skills_text, re.IGNORECASE):
                    add_if_found(raw_tech, skill)
                continue

            # short tokens (1-3 chars) require exact token match
            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_tech, skill)
                continue

            # otherwise use whole-word boundary check
            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_tech, skill)

        for skill in soft_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, combined_skills_text, re.IGNORECASE):
                    add_if_found(raw_soft, skill)
                continue

            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_soft, skill)
                continue

            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_soft, skill)

        # 3) Compact-text fallback only for multi-word skills (avoid false positives for short tokens)
        compact_text = re.sub(r'[^a-z0-9]+', '', text_lower)
        for skill in tech_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_tech, skill):
                    raw_tech.append(skill)
        for skill in soft_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_soft, skill):
                    raw_soft.append(skill)

        # 4) Normalize result to canonical casing from master lists and dedupe preserving order
        def canonical_map(master_list):
            cm = {}
            for s in master_list:
                if isinstance(s, str) and s.strip():
                    cm[s.strip().lower()] = s.strip()
            return cm

        tech_canon = canonical_map(tech_master)
        soft_canon = canonical_map(soft_master)

        def normalize_and_dedupe(raw_list, canon_map):
            out = []
            seen = set()
            for s in raw_list:
                key = s.strip().lower()
                if not key or key in seen:
                    continue
                seen.add(key)
                out.append(canon_map.get(key, s.strip().title()))
            return out

        technical = normalize_and_dedupe(raw_tech, tech_canon)
        soft = normalize_and_dedupe(raw_soft, soft_canon)

        # Return as comma-separated strings
        technical_str = ", ".join(technical) if technical else ""
        soft_str = ", ".join(soft) if soft else ""
        
        return technical_str, soft_str

    def extract_skills(self, text, sections=None):
        """
        High-precision skill extraction that prioritizes skills section:
        - First searches in dedicated skills section if available
        - Falls back to entire resume if no skills section found
        - Uses PhraseMatcher for high precision multi-word skills
        - Exact token matching for short/special tokens (C, R, Go, SQL, etc.)
        - Whole-word regex for skills with punctuation (C++, C#, node.js)
        - Normalize and dedupe while preserving discovery order
        Returns: (technical_list, soft_list)
        """
        import re

        if not text:
            return [], []

        # Determine search area: prioritize skills section if available
        search_text = text
        search_area_label = "entire resume"
        
        if sections and isinstance(sections, dict):
            skills_section = sections.get('skills', '').strip()
            if skills_section and len(skills_section) > 20:
                search_text = skills_section
                search_area_label = "skills section"
        
        # Create NLP document for the search area
        doc = self.nlp(search_text)
        text_lower = search_text.lower()

        raw_tech = []
        raw_soft = []

        # 1) PhraseMatcher results (high precision)
        for _id, start, end in self.tech_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_tech.append(span)
        for _id, start, end in self.soft_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_soft.append(span)

        # token set for exact matches
        token_texts = [t.text for t in doc if t.text.strip()]
        token_set = {t.lower() for t in token_texts}

        # Flatten master lists
        def flatten_master(master):
            out = []
            if isinstance(master, dict):
                for _, lst in master.items():
                    out.extend(lst)
            elif isinstance(master, list):
                out.extend(master)
            return out

        tech_master = flatten_master(TECHNICAL_SKILLS)
        soft_master = flatten_master(SOFT_SKILLS)

        # helpers
        def already_found(raw_list, skill):
            return skill.strip().lower() in {s.strip().lower() for s in raw_list}

        def add_if_found(raw_list, skill):
            if not already_found(raw_list, skill):
                raw_list.append(skill)

        # 2) Exact token matching for short/special skills (prevent substring noise)
        for skill in tech_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            # skills containing punctuation (C++, C#, node.js) -> regex whole-word check
            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, search_text, re.IGNORECASE):
                    add_if_found(raw_tech, skill)
                continue

            # short tokens (1-3 chars) require exact token match
            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_tech, skill)
                continue

            # otherwise use whole-word boundary check
            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_tech, skill)

        for skill in soft_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, search_text, re.IGNORECASE):
                    add_if_found(raw_soft, skill)
                continue

            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_soft, skill)
                continue

            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_soft, skill)

        # 3) Compact-text fallback only for multi-word skills (avoid false positives for short tokens)
        compact_text = re.sub(r'[^a-z0-9]+', '', text_lower)
        for skill in tech_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_tech, skill):
                    raw_tech.append(skill)
        for skill in soft_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_soft, skill):
                    raw_soft.append(skill)

        # 4) Normalize result to canonical casing from master lists and dedupe preserving order
        def canonical_map(master_list):
            cm = {}
            for s in master_list:
                if isinstance(s, str) and s.strip():
                    cm[s.strip().lower()] = s.strip()
            return cm

        tech_canon = canonical_map(tech_master)
        soft_canon = canonical_map(soft_master)

        def normalize_and_dedupe(raw_list, canon_map):
            out = []
            seen = set()
            for s in raw_list:
                key = s.strip().lower()
                if not key or key in seen:
                    continue
                seen.add(key)
                out.append(canon_map.get(key, s.strip().title()))
            return out

        technical = normalize_and_dedupe(raw_tech, tech_canon)
        soft = normalize_and_dedupe(raw_soft, soft_canon)

        return technical, soft

    # education extraction removed per user request
        
    def extract_personal_information(self, text, sections):
        """Extract personal information fields exactly as they appear in the resume"""
        import re
        
        personal_info = {
            'name': '',
            'email': '',
            'phone_10_digits': '',
            'linkedin': '',
            'github': '',
            'leetcode': ''
        }

        # Extract name following strict priority rules described by the user.
        # 1) Topmost non-empty text block before contact info keywords
        # 2) Apply capitalization/word rules (2-4 words, no digits, allow initials and hyphens)
        lines = [ln.strip() for ln in text.split('\n')]
        # look only at top portion of document (first 15-20 lines)
        header_lines = lines[:20]

        contact_keywords = ('email', 'phone', 'mobile', 'linkedin', 'portfolio', 'contact', 'github', 'leetcode')
        stop_index = None
        for i, ln in enumerate(header_lines):
            if not ln:
                continue
            low = ln.lower()
            # stop searching when we reach an obvious contact line
            if any(kw in low for kw in contact_keywords) or '@' in ln or re.search(r'\d{3,}', ln):
                stop_index = i
                break

        search_lines = header_lines[:stop_index] if stop_index is not None else header_lines

        # helper regexes
        word_pat = re.compile(r"^(?:[A-Z][a-zA-Z]*(?:[-'][A-Za-z][a-zA-Z]*)*|[A-Z]\.?)(?:\.)?$")
        banned_words = set(['resume', 'curriculum', 'vitae', 'cv', 'profile', 'objective', 'engineer', 'developer', 'intern', 'student'])
        org_suffixes = ('ltd', 'inc', 'university', 'college', 'institute', 'company', 'corp', 'pvt', 'llc')

        def looks_like_name(s: str) -> bool:
            if not s or len(s) < 2 or len(s) > 60:
                return False
            if any(ch.isdigit() for ch in s):
                return False
            low = s.lower()
            if any(b in low for b in banned_words):
                return False
            if any(suffix in low for suffix in org_suffixes):
                return False
            words = [w.strip(".,") for w in s.split() if w.strip(".,")]
            if not (2 <= len(words) <= 4):
                return False
            for w in words:
                if not word_pat.match(w):
                    return False
            return True

        # 1) pick first matching topmost line
        for ln in search_lines:
            if not ln:
                continue
            if looks_like_name(ln):
                personal_info['name'] = ln
                break

        # 2) Fallback: if we found a contact line, pick the nearest non-empty line immediately above it
        if not personal_info['name'] and stop_index is not None and stop_index > 0:
            for j in range(stop_index - 1, max(-1, stop_index - 4), -1):
                cand = header_lines[j].strip()
                if cand and looks_like_name(cand):
                    personal_info['name'] = cand
                    break
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            personal_info['email'] = email_match.group()
        
        # Extract phone (special rule: 10 digits only)
        phone_patterns = [
            r'\+?[\d\s\-\(\)]{10,}',  # General phone pattern
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # US format
            r'\b\d{10}\b',  # 10 digits
            r'\+91[-.\s]?\d{5}[-.\s]?\d{5}',  # Indian format
        ]
        
        for pattern in phone_patterns:
            phone_matches = re.findall(pattern, text)
            for phone in phone_matches:
                # Extract only digits
                digits = re.sub(r'\D', '', phone)
                if len(digits) >= 10:
                    personal_info['phone_10_digits'] = digits[-10:]  # Keep last 10 digits
                    break
            if personal_info['phone_10_digits']:
                break
        
        # Extract LinkedIn - look for properly formatted LinkedIn URLs
        linkedin_patterns = [
            r'(?:linkedin\.com/in|linkedin\.com/profile|linkedin\.com/pub)/[\w\-]+(?:/[^\s]*)?',
            r'https?://(?:www\.)?linkedin\.com/in/[\w\-]+(?:/[^\s]*)?'
        ]
        
        # First try finding LinkedIn near the word "LinkedIn" or "Profile"
        context_pattern = r'(?:linkedin|profile|contact).*?((?:linkedin\.com|(?:www\.)?linkedin\.com)/\S+)'
        context_match = re.search(context_pattern, text.lower(), re.IGNORECASE | re.DOTALL)
        if context_match:
            url = context_match.group(1)
            if not url.endswith(('.pdf', '.doc', '.html')):  # Basic URL cleanup
                personal_info['linkedin'] = url.strip('."\')')
        else:
            # Fall back to pattern matching if no contextual match
            for pattern in linkedin_patterns:
                linkedin_match = re.search(pattern, text, re.IGNORECASE)
                if linkedin_match:
                    url = linkedin_match.group()
                    if not url.endswith(('.pdf', '.doc', '.html')):
                        personal_info['linkedin'] = url.strip('."\')')
                        break
        
        # Extract GitHub - look for properly formatted GitHub URLs
        github_patterns = [
            r'github\.com/[\w\-]+',
            r'https?://(?:www\.)?github\.com/[\w\-]+'
        ]
        
        # First look near the word "GitHub"
        github_context = r'(?:github|repo).*?((?:github\.com|(?:www\.)?github\.com)/\S+)'
        github_context_match = re.search(github_context, text.lower(), re.IGNORECASE | re.DOTALL)
        if github_context_match:
            url = github_context_match.group(1)
            if not url.endswith(('.pdf', '.doc', '.html')):
                personal_info['github'] = url.strip('."\')')
        else:
            for pattern in github_patterns:
                github_match = re.search(pattern, text, re.IGNORECASE)
                if github_match:
                    url = github_match.group()
                    if not url.endswith(('.pdf', '.doc', '.html')):
                        personal_info['github'] = url.strip('."\')')
                        break
        
        # Extract LeetCode - only extract if explicitly labeled as LeetCode
        leetcode_patterns = [
            r'leetcode\.com/[\w\-]+',
            r'https?://(?:www\.)?leetcode\.com/[\w\-]+'
        ]
        
        # Only match LeetCode if the word "LeetCode" appears nearby
        leetcode_context = r'leetcode.*?((?:leetcode\.com|(?:www\.)?leetcode\.com)/\S+)'
        leetcode_match = re.search(leetcode_context, text.lower(), re.IGNORECASE | re.DOTALL)
        if leetcode_match:
            url = leetcode_match.group(1)
            if not url.endswith(('.pdf', '.doc', '.html')):
                personal_info['leetcode'] = url.strip('."\')')

        return personal_info


# ...existing code...
    # ...existing code...
           
# ...existing code...
    def keyword_matching(self, text, job_role):
        if job_role not in JOB_KEYWORDS:
            return [], 0
        keywords = JOB_KEYWORDS[job_role]
        text_lower = text.lower()
        found_keywords = []
        for keyword in keywords:
            if keyword.lower() in text_lower:
                found_keywords.append(keyword)
        match_percentage = (len(found_keywords) / len(keywords)) * 100
        return found_keywords, match_percentage


    def calculate_ats_score(self, text, sections):
        score = 0
        required_sections = ['experience', 'education', 'skills']
        for section in required_sections:
            if sections.get(section) and len(sections[section]) > 50:
                score += 13.33
        word_count = len(text.split())
        if 300 <= word_count <= 800:
            score += 20
        elif word_count > 200:
            score += 10
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        if re.search(email_pattern, text):
            score += 10
        if re.search(phone_pattern, text):
            score += 10
        bullet_patterns = [r'•', r'◦', r'\*', r'-\s', r'→']
        bullet_count = sum(len(re.findall(pattern, text)) for pattern in bullet_patterns)
        if bullet_count >= 5:
            score += 20
        elif bullet_count >= 2:
            score += 10
        return min(score, 100)


    def analyze_resume(self, text, job_role):
        sections = self.extract_sections(text)
        tech_skills, soft_skills = self.extract_skills(text)
        found_keywords, match_percentage = self.keyword_matching(text, job_role)
        ats_score = self.calculate_ats_score(text, sections)
        overall_score = (ats_score + match_percentage) / 2
       
        analysis_data = {
            "sections": sections,
            "technical_skills": tech_skills,
            "soft_skills": soft_skills,
            "found_keywords": found_keywords,
            "missing_keywords": [kw for kw in JOB_KEYWORDS[job_role] if kw not in found_keywords],
            "role_match_percentage": match_percentage,
            "ats_score": ats_score,
            "overall_score": overall_score,
            "word_count": len(text.split()),
            "section_count": len([s for s in sections.values() if s])
        }
       
        # Set context for AI assistant
        # self.ai_assistant.set_resume_context(analysis_data) # AI Assistant removed
       
        return analysis_data

    def extract_skills_from_sections_only(self, text, sections=None):
        """
        Extract skills ONLY from dedicated 'Skills', 'Technical Skills', or 'Soft Skills' sections.
        Does not extract from project descriptions, experience, or summary sections.
        Returns: (technical_skills, soft_skills) as comma-separated strings
        """
        import re

        if not text:
            return "", ""

        # Get sections if not provided
        if sections is None:
            sections = self.extract_sections(text)

        # Look for skills sections with various naming patterns
        skills_sections = []
        
        # Check for different section names
        section_patterns = {
            'skills': r'(skills|technical skills|soft skills|competencies|expertise|abilities|technologies)',
            'technical_skills': r'(technical skills|technical competencies|programming|languages|tools|frameworks)',
            'soft_skills': r'(soft skills|interpersonal skills|leadership skills|communication skills|management skills)'
        }
        
        # Extract all skills-related sections
        for section_name, pattern in section_patterns.items():
            if section_name in sections and sections[section_name].strip():
                skills_sections.append(sections[section_name].strip())
        
        # If no dedicated skills sections found, return empty
        if not skills_sections:
            return "", ""
        
        # Combine all skills sections
        combined_skills_text = "\n".join(skills_sections)
        
        # Create NLP document for the skills sections only
        doc = self.nlp(combined_skills_text)
        text_lower = combined_skills_text.lower()

        raw_tech = []
        raw_soft = []

        # 1) PhraseMatcher results (high precision)
        for _id, start, end in self.tech_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_tech.append(span)
        for _id, start, end in self.soft_matcher(doc):
            span = doc[start:end].text.strip()
            if span:
                raw_soft.append(span)

        # token set for exact matches
        token_texts = [t.text for t in doc if t.text.strip()]
        token_set = {t.lower() for t in token_texts}

        # Flatten master lists
        def flatten_master(master):
            out = []
            if isinstance(master, dict):
                for _, lst in master.items():
                    out.extend(lst)
            elif isinstance(master, list):
                out.extend(master)
            return out

        tech_master = flatten_master(TECHNICAL_SKILLS)
        soft_master = flatten_master(SOFT_SKILLS)

        # helpers
        def already_found(raw_list, skill):
            return skill.strip().lower() in {s.strip().lower() for s in raw_list}

        def add_if_found(raw_list, skill):
            if not already_found(raw_list, skill):
                raw_list.append(skill)

        # 2) Exact token matching for short/special skills (prevent substring noise)
        for skill in tech_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            # skills containing punctuation (C++, C#, node.js) -> regex whole-word check
            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, combined_skills_text, re.IGNORECASE):
                    add_if_found(raw_tech, skill)
                continue

            # short tokens (1-3 chars) require exact token match
            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_tech, skill)
                continue

            # otherwise use whole-word boundary check
            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_tech, skill)

        for skill in soft_master:
            if not isinstance(skill, str) or not skill.strip():
                continue
            key = skill.strip()
            key_lower = key.lower()

            if re.search(r'[^A-Za-z0-9\s]', key):
                pattern = r'(?<!\w)' + re.escape(key) + r'(?!\w)'
                if re.search(pattern, combined_skills_text, re.IGNORECASE):
                    add_if_found(raw_soft, skill)
                continue

            if len(re.sub(r'\s+', '', key_lower)) <= 3:
                if key_lower in token_set:
                    add_if_found(raw_soft, skill)
                continue

            pattern = r'(?<!\w)' + re.escape(key_lower) + r'(?!\w)'
            if re.search(pattern, text_lower):
                add_if_found(raw_soft, skill)

        # 3) Compact-text fallback only for multi-word skills (avoid false positives for short tokens)
        compact_text = re.sub(r'[^a-z0-9]+', '', text_lower)
        for skill in tech_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_tech, skill):
                    raw_tech.append(skill)
        for skill in soft_master:
            if not isinstance(skill, str):
                continue
            key = skill.strip().lower()
            if ' ' in key and len(re.sub(r'[^a-z0-9]+', '', key)) >= 6:
                compact_skill = re.sub(r'[^a-z0-9]+', '', key)
                if compact_skill and compact_skill in compact_text and not already_found(raw_soft, skill):
                    raw_soft.append(skill)

        # 4) Normalize result to canonical casing from master lists and dedupe preserving order
        def canonical_map(master_list):
            cm = {}
            for s in master_list:
                if isinstance(s, str) and s.strip():
                    cm[s.strip().lower()] = s.strip()
            return cm

        tech_canon = canonical_map(tech_master)
        soft_canon = canonical_map(soft_master)

        def normalize_and_dedupe(raw_list, canon_map):
            out = []
            seen = set()
            for s in raw_list:
                key = s.strip().lower()
                if not key or key in seen:
                    continue
                seen.add(key)
                out.append(canon_map.get(key, s.strip().title()))
            return out

        technical = normalize_and_dedupe(raw_tech, tech_canon)
        soft = normalize_and_dedupe(raw_soft, soft_canon)

        # Return as comma-separated strings
        technical_str = ", ".join(technical) if technical else ""
        soft_str = ", ".join(soft) if soft else ""
        
        return technical_str, soft_str


# (AI Assistant removed as per request)


# ---------------------------
# PDF Generation Functions
# ---------------------------
def generate_resume_report_pdf(analysis_data, filename, job_role):
    """Generate a comprehensive PDF report of the resume analysis"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
   
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.HexColor('#1e40af')
    )
   
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        textColor=colors.HexColor('#1e40af')
    )
   
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
   
    # Build the PDF content
    story = []
   
    # Title
    story.append(Paragraph("Resume Analysis Report", title_style))
    story.append(Spacer(1, 20))
   
    # File information
    story.append(Paragraph(f"<b>File:</b> {filename}", normal_style))
    story.append(Paragraph(f"<b>Target Role:</b> {job_role}", normal_style))
    story.append(Paragraph(f"<b>Analysis Date:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    story.append(Spacer(1, 20))
   
    # Summary scores
    story.append(Paragraph("Overall Scores", heading_style))
   
    scores_data = [
        ['Metric', 'Score', 'Status'],
        ['ATS Compatibility', f"{analysis_data.get('ats_score', 0):.1f}/100",
         'Excellent' if analysis_data.get('ats_score', 0) >= 80 else
         'Good' if analysis_data.get('ats_score', 0) >= 60 else 'Needs Improvement'],
        ['Role Match', f"{analysis_data.get('role_match_percentage', 0):.1f}%",
         'Excellent' if analysis_data.get('role_match_percentage', 0) >= 70 else
         'Good' if analysis_data.get('role_match_percentage', 0) >= 50 else 'Needs Improvement'],
        ['Overall Score', f"{analysis_data.get('overall_score', 0):.1f}/100",
         'Excellent' if analysis_data.get('overall_score', 0) >= 80 else
         'Good' if analysis_data.get('overall_score', 0) >= 60 else 'Needs Improvement']
    ]
   
    scores_table = Table(scores_data, colWidths=[2*inch, 1.5*inch, 2*inch])
    scores_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
   
    story.append(scores_table)
    story.append(Spacer(1, 20))

    # Technical Skills
    story.append(Paragraph("Technical Skills Detected", heading_style))
    tech_skills = analysis_data.get('technical_skills', [])
    if tech_skills:
        tech_text = ", ".join(tech_skills)
        story.append(Paragraph(f"<b>Count:</b> {len(tech_skills)} skills", normal_style))
        story.append(Paragraph(f"<b>Skills:</b> {tech_text}", normal_style))
    else:
        story.append(Paragraph("No technical skills detected", normal_style))
    story.append(Spacer(1, 15))
   
    # Soft Skills
    story.append(Paragraph("Soft Skills Detected", heading_style))
    soft_skills = analysis_data.get('soft_skills', [])
    if soft_skills:
        soft_text = ", ".join(soft_skills)
        story.append(Paragraph(f"<b>Count:</b> {len(soft_skills)} skills", normal_style))
        story.append(Paragraph(f"<b>Skills:</b> {soft_text}", normal_style))
    else:
        story.append(Paragraph("No soft skills detected", normal_style))
    story.append(Spacer(1, 15))
   
    # Keywords Analysis
    story.append(Paragraph("Keywords Analysis", heading_style))
    found_keywords = analysis_data.get('found_keywords', [])
    missing_keywords = analysis_data.get('missing_keywords', [])
   
    story.append(Paragraph(f"<b>Found Keywords:</b> {len(found_keywords)}", normal_style))
    if found_keywords:
        found_text = ", ".join(found_keywords[:20])  # Limit to first 20
        story.append(Paragraph(f"Keywords: {found_text}", normal_style))
   
    story.append(Paragraph(f"<b>Missing Keywords:</b> {len(missing_keywords)}", normal_style))
    if missing_keywords:
        missing_text = ", ".join(missing_keywords[:20])  # Limit to first 20
        story.append(Paragraph(f"Suggested: {missing_text}", normal_style))
   
    story.append(Spacer(1, 15))
   
    # Recommendations removed to align with simplified UI
    story.append(Spacer(1, 20))
   
    # Footer
    story.append(Paragraph("Generated by Resume Insight Dashboard",
                          ParagraphStyle('Footer', parent=styles['Normal'],
                                        fontSize=8, alignment=1,
                                        textColor=colors.grey)))
   
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


# ---------------------------
# Helper UI functions
# ---------------------------
def display_metric_card(title, value, description=""):
    st.markdown(f"""
    <div class="metric-card fade-in-up">
        <div class="metric-value">{value}</div>
        <div class="metric-label">{title}</div>
        {f"<small>{description}</small>" if description else ""}
    </div>
    """, unsafe_allow_html=True)


def display_alert_box(message, alert_type="info"):
    st.markdown(f"""
    <div class="alert-modern alert-{alert_type} fade-in-up">
        {message}
    </div>
    """, unsafe_allow_html=True)


# Render AI Assistant messages with numbering and topic-based color classes
def render_ai_message(message):
    """
    Formats AI assistant message with clean, properly aligned formatting:
    - Displays clean numbered points
    - Removes markdown formatting
    - Ensures proper alignment
    - Chooses CSS class based on detected topic keywords
    """
    # Clean the message of any remaining markdown
    import re
    
    # Remove any remaining markdown formatting
    clean_msg = re.sub(r'\*\*(.*?)\*\*', r'\1', message)
    clean_msg = re.sub(r'\*(.*?)\*', r'\1', clean_msg)
    clean_msg = re.sub(r'`(.*?)`', r'\1', clean_msg)
    
    # Split into lines and process
    lines = [ln.strip() for ln in clean_msg.split("\n") if ln.strip()]
    
    # Format lines with proper numbering and alignment
    formatted_lines = []
    for line in lines:
        # If line already starts with a number, keep it as is
        if re.match(r'^\d+\.', line):
            formatted_lines.append(line)
        # If line starts with common prefixes, format it
        elif line.startswith(('To enhance', 'To improve', 'To create', 'To make')):
            formatted_lines.append(f"1. {line}")
        else:
            # Add numbering if it doesn't have it
            if not re.match(r'^\d+\.', line):
                formatted_lines.append(f"• {line}")
            else:
                formatted_lines.append(line)
    
    # Join with proper line breaks
    formatted = "<br>".join(formatted_lines)
    
    # Detect topic for CSS coloring
    lower = clean_msg.lower()
    if 'skill' in lower or 'skills' in lower or 'skills section' in lower:
        cls = "chat-ai-skills"
    elif 'ats' in lower or 'applicant tracking' in lower or 'ats-friendly' in lower:
        cls = "chat-ai-ats"
    elif 'experience' in lower or 'work history' in lower or 'experience section' in lower:
        cls = "chat-ai-experience"
    elif 'keyword' in lower or 'keywords' in lower or 'keyword strategy' in lower:
        cls = "chat-ai-keywords"
    else:
        cls = "chat-ai"

    st.markdown(f"""
    <div class="chat-message {cls} fade-in-up">
        <strong>🤖 AI Assistant:</strong>
        <div style="margin-top:0.5rem; line-height: 1.6;">{formatted}</div>
    </div>
    """, unsafe_allow_html=True)


# ---------------------------
# Login / Dashboard UI (main app) - with updated chat rendering
# ---------------------------
def login_page():
    st.markdown("""
    <div class="login-container fade-in-up">
        <h2 style="text-align: center; color: #1e293b; margin-bottom: 2rem; font-weight: 700;">
            📊 RESUME INSIGHT
        </h2>
        <p style="text-align: center; color: #64748b; margin-bottom: 2rem;">
            Resume Analysis Dashboard
        </p>
    </div>
    """, unsafe_allow_html=True)


    if 'show_register' not in st.session_state:
        st.session_state.show_register = False


    if not st.session_state.show_register:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("### 🔐 Login to Your Account")
            with st.form("login_form"):
                username = st.text_input("👤 Username")
                password = st.text_input("🔒 Password", type="password")
                login_button = st.form_submit_button("🚀 Login", use_container_width=True)
                if login_button:
                    db = DatabaseManager()
                    user = db.authenticate_user(username, password)
                    if user:
                        st.session_state.user = user
                        st.session_state.logged_in = True
                        st.success("🎉 Login successful!")
                        st.rerun()
                    else:
                        st.error("❌ Invalid username or password")
            if st.button("🆕 New user? Create account", use_container_width=True):
                st.session_state.show_register = True
                st.rerun()
    else:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("### ✨ Create New Account")
            with st.form("register_form"):
                new_username = st.text_input("👤 New Username")
                new_email = st.text_input("📧 Email")
                new_password = st.text_input("🔒 New Password", type="password")
                confirm_password = st.text_input("🔒 Confirm Password", type="password")
                register_button = st.form_submit_button("🚀 Create Account", use_container_width=True)
                if register_button:
                    if new_password != confirm_password:
                        st.error("❌ Passwords do not match")
                    elif len(new_password) < 6:
                        st.error("❌ Password must be at least 6 characters long")
                    else:
                        db = DatabaseManager()
                        if db.register_user(new_username, new_email, new_password):
                            st.success("🎉 Registration successful! Please login.")
                            st.session_state.show_register = False
                            st.rerun()
                        else:
                            st.error("❌ Registration failed. Username or email may already exist.")
            if st.button("↩️ Back to Login", use_container_width=True):
                st.session_state.show_register = False
                st.rerun()


def dashboard_page():
    user = st.session_state.user
    st.markdown(f"""
    <div class="modern-header fade-in-up">
        <div class="header-title">Welcome back, {user['username']}! 👋</div>
        <div class="header-subtitle">📊 Resume Analysis Dashboard</div>
    </div>
    """, unsafe_allow_html=True)


    with st.sidebar:
        st.markdown("""
        <div class="sidebar-modern fade-in-up">
            <h3 style="color: #334155; margin-bottom: 1rem;">⚙️ Analysis Configuration</h3>
        </div>
        """, unsafe_allow_html=True)


        # Job role selection with auto-detect option
        job_roles = ["Auto-detect"] + list(JOB_KEYWORDS.keys())
        selected_role_option = st.selectbox("🎯 Target Job Role:", job_roles, help="Choose a role or select Auto-detect")


        st.markdown("---")
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.user = None
            st.rerun()


        st.markdown("""
        <div class="sidebar-modern fade-in-up">
            <h3 style="color: #334155; margin-bottom: 1rem;">📊 Analysis History</h3>
        </div>
        """, unsafe_allow_html=True)


        db = DatabaseManager()
        analyses = db.get_user_analyses(user['id'])
        if analyses:
            for analysis in analyses[:5]:
                with st.expander(f"📄 {analysis['filename']} - {analysis['job_role']}"):
                    st.markdown(f"""
                    <div class="history-card">
                        <p><strong>📅 Date:</strong> {analysis['created_at'].strftime('%Y-%m-%d %H:%M')}</p>
                        <p><strong>🎯 ATS Score:</strong> {analysis['ats_score']}/100</p>
                        <p><strong>📈 Role Match:</strong> {analysis['role_match_percentage']:.1f}%</p>
                        <p><strong>⭐ Overall Score:</strong> {analysis['overall_score']:.1f}/100</p>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("📝 No previous analyses found")


    st.markdown('<div class="section-header-modern fade-in-up">📄 Upload Your Resume</div>', unsafe_allow_html=True)


    uploaded_file = st.file_uploader("📁 Select your resume file", type=['pdf', 'docx', 'txt'], help="Supported formats: PDF, DOCX, TXT")


    if uploaded_file is not None:
        file_details = {"📄 Filename": uploaded_file.name, "📊 File size": f"{uploaded_file.size / 1024:.1f} KB", "🔧 File type": uploaded_file.type}
        with st.expander("📋 File Information", expanded=False):
            for key, value in file_details.items():
                st.write(f"**{key}:** {value}")


        # Initialize analyzer with environment API key
        analyzer = ResumeAnalyzer()
        
        with st.spinner("🔄 Processing your resume..."):
            try:
                if uploaded_file.type == "application/pdf":
                    text = analyzer.extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text = analyzer.extract_text_from_docx(uploaded_file)
                else:
                    text = analyzer.extract_text_from_txt(uploaded_file)
            except Exception as e:
                st.error(f"❌ Error extracting text: {str(e)}")
                return


        if "Error" not in text and text.strip():
            st.success("✅ Resume processed successfully!")
            
            # Build RAG index for resume question answering
            with st.spinner("🔍 Building knowledge base for AI Assistant..."):
                rag_built, rag_error = analyzer.build_resume_rag_index(text)
                if rag_built:
                    st.success("✅ AI Assistant ready to answer questions about this resume!")
                elif rag_error == "quota_exceeded":
                    st.error("""
                    🚫 **OpenAI Quota Exceeded**
                    
                    Your OpenAI API key has run out of credits. LangChain RAG AI Assistant is unavailable.
                    
                    **To fix this:**
                    1. Visit https://platform.openai.com/settings/organization/billing
                    2. Add $10 credits (analyze 200-500 resumes)
                    3. Alternatively, use a different API key with available credits
                    
                    💡 **Good news:** Other resume analysis features still work:
                    - ATS scoring
                    - Skills detection  
                    - Keyword analysis
                    - PDF report generation
                    """)
                elif rag_error == "invalid_key":
                    st.error("""
                    🔑 **Invalid OpenAI API Key**
                    
                    Your API key is invalid or expired.
                    
                    **To fix this:**
                    1. Get a valid key from https://platform.openai.com/api-keys
                    2. Update your 'env' file: OPENAI_API_KEY=sk-proj-...
                    3. Restart the application
                    """)
                else:
                    st.warning(f"⚠️ AI Assistant unavailable: {rag_error}. Other features will work normally.")


            # Determine target role: manual selection or auto-detect
            if selected_role_option != "Auto-detect":
                detected_role = selected_role_option
            else:
                detected_role = None
                best_score = -1
                text_lower = text.lower()
                for role, keywords in JOB_KEYWORDS.items():
                    score = sum(1 for kw in keywords if kw.lower() in text_lower)
                    if score > best_score:
                        best_score = score
                        detected_role = role


            if selected_role_option == "Auto-detect":
                if detected_role:
                    st.info(f"🎯 Detected target role: {detected_role}")
                else:
                    detected_role = "Software Engineer"
                    st.warning("⚠️ Could not auto-detect a role. Defaulting to Software Engineer.")


            try:
                results = analyzer.analyze_resume(text, detected_role)
                # Defensive checks: ensure results is a dict and sections is a dict
                if results is None:
                    raise ValueError("Analyzer returned no results. The resume text may be malformed or unsupported format.")
                if not isinstance(results, dict):
                    raise ValueError(f"Analyzer returned unexpected type: {type(results)}")
                # Normalize sections to an empty dict if missing/None to avoid subscript errors
                sections = results.get('sections')
                if sections is None or not isinstance(sections, dict):
                    sections = {}
                    results['sections'] = sections
                db = DatabaseManager()
                if db.save_analysis(user['id'], uploaded_file.name, detected_role, results, text):
                    st.success("💾 Analysis saved to history!")
                else:
                    st.warning("⚠️ Could not save analysis to history")


                # Add PDF download button
                st.markdown("---")
                col1, col2, col3 = st.columns([1, 1, 1])
                with col2:
                    if st.button("📄 Download Analysis Report (PDF)", use_container_width=True, type="primary"):
                        pdf_buffer = generate_resume_report_pdf(results, uploaded_file.name, detected_role)
                        st.download_button(
                            label="📥 Download PDF Report",
                            data=pdf_buffer.getvalue(),
                            file_name=f"resume_analysis_{uploaded_file.name}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )


                tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Personal Information", "📊 Summary", "🔍 Skills Analysis", "🎯 ATS Analysis", "🤖 AI Assistant"])


                # ----- TAB 1: Personal Information -----
                with tab1:
                    st.markdown('<div class="section-header-modern fade-in-up">📋 Personal Information</div>', unsafe_allow_html=True)
                    
                    # Extract personal information (use safe .get for sections)
                    personal_info = analyzer.extract_personal_information(text, results.get('sections', {}))
                    
                    # Display personal information in a structured format
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("### 👤 Basic Information")
                        
                        # Name
                        if personal_info['name']:
                            st.markdown(f"**Name:** {personal_info['name']}")
                        else:
                            st.markdown("**Name:** Not included")
                        
                        # Email
                        if personal_info['email']:
                            st.markdown(f"**Email:** {personal_info['email']}")
                        else:
                            st.markdown("**Email:** Not included")
                        
                        # Phone
                        if personal_info['phone_10_digits']:
                            st.markdown(f"**Phone:** {personal_info['phone_10_digits']}")
                        else:
                            st.markdown("**Phone:** Not included")
                        
                        # LinkedIn
                        if personal_info['linkedin']:
                            st.markdown(f"**LinkedIn:** {personal_info['linkedin']}")
                        else:
                            st.markdown("**LinkedIn:** Not included")
                        
                        # GitHub
                        if personal_info['github']:
                            st.markdown(f"**GitHub:** {personal_info['github']}")
                        else:
                            st.markdown("**GitHub:** Not included")
                        
                        # LeetCode - only show when present
                        if personal_info.get('leetcode'):
                            st.markdown(f"**LeetCode:** {personal_info['leetcode']}")
                    
                    with col2:
                        # Education display intentionally hidden
                        pass


                # ----- TAB 2: Summary -----
                with tab2:
                    st.markdown('<div class="section-header-modern fade-in-up">📊 Resume Summary</div>', unsafe_allow_html=True)
                    metric_cols = st.columns(3)
                    with metric_cols[0]:
                        ats_val = results.get('ats_score', 0) or 0
                        display_metric_card("ATS Score", f"{ats_val:.1f}/100")
                    with metric_cols[1]:
                        role_match_val = results.get('role_match_percentage', 0) or 0
                        display_metric_card("Role Match", f"{role_match_val:.1f}%")
                    with metric_cols[2]:
                        word_count_val = results.get('word_count', 0) or 0
                        display_metric_card("Word Count", f"{word_count_val}")


                    overall_score = results.get('overall_score', 0) or 0
                    if overall_score >= 80:
                        display_alert_box("🎉 Excellent resume! Your resume shows strong alignment with the target role and good ATS compatibility.", "success")
                    elif overall_score >= 60:
                        display_alert_box("⚠️ Good foundation with room for improvement. Focus on adding more role-specific keywords and optimizing for ATS.", "warning")
                    else:
                        display_alert_box("🚨 Significant improvements needed. Consider restructuring sections, adding relevant keywords, and improving ATS compatibility.", "error")


                # ----- TAB 3: Skills Analysis -----
                with tab3:
                    st.markdown('<div class="section-header-modern fade-in-up">🔍 Skills Analysis</div>', unsafe_allow_html=True)
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("### 💻 Technical Skills Detected")
                        tech_skills = results.get('technical_skills', []) or []
                        if tech_skills:
                            tech_html = "".join([f'<span class="skill-tag">{skill}</span>' for skill in tech_skills])
                            st.markdown(tech_html, unsafe_allow_html=True)
                            st.metric("Technical Skills Count", len(tech_skills))
                        else:
                            display_alert_box("⚠️ No technical skills detected. Consider adding a dedicated skills section.", "warning")
                    with col2:
                        st.markdown("### 🤝 Soft Skills Detected")
                        soft_skills = results.get('soft_skills', []) or []
                        if soft_skills:
                            soft_html = "".join([f'<span class="skill-tag">{skill}</span>' for skill in soft_skills])
                            st.markdown(soft_html, unsafe_allow_html=True)
                            st.metric("Soft Skills Count", len(soft_skills))
                        else:
                            display_alert_box("ℹ️ Limited soft skills detected. Consider highlighting leadership, communication, and teamwork skills.", "info")


                    st.divider()
                    st.markdown("### 🎯 Analysis for " + detected_role)
                    progress_col, details_col = st.columns([1, 2])
                    with progress_col:
                        match_percentage = results.get('role_match_percentage', 0) or 0
                        st.metric("Match Percentage", f"{match_percentage:.1f}%")
                        st.progress(match_percentage / 100)
                    with details_col:
                        if match_percentage >= 70:
                            display_alert_box("🎉 Excellent match for this role! Your skills align well with industry expectations.", "success")
                        elif match_percentage >= 50:
                            display_alert_box("⚠️ Good match with opportunities for improvement. Consider adding more role-specific skills.", "warning")
                        else:
                            display_alert_box("🚨 Limited match detected. Focus on adding more relevant skills and keywords for this role.", "error")


                    # Suggested Keywords (clean chips + Pro Tip)
                    missing_keywords = results.get('missing_keywords', []) or []
                    if missing_keywords:
                        st.markdown(f"""
                        <div class="section-header-modern fade-in-up" style="margin-top: 1.5rem;">🔑 Suggested Keywords to Add</div>
                        <p style="color: #475569; margin-bottom: 1rem; font-size: 1rem;">
                            💡 These keywords are commonly found in <strong>{detected_role}</strong> job descriptions:
                        </p>
                        """, unsafe_allow_html=True)
                        missing_html = "".join([f'<span class="keyword-suggested">{kw}</span>' for kw in missing_keywords[:30]])
                        st.markdown(missing_html, unsafe_allow_html=True)
                        st.markdown("""
                        <div class="alert-modern alert-info fade-in-up">
                            💡 <strong>Pro Tip:</strong> Try to naturally incorporate these keywords into your resume sections.
                            Don't just add them randomly — weave them into your experience descriptions and skills sections.
                        </div>
                        """, unsafe_allow_html=True)



                    # Content Analysis and Tone Category
                    st.divider()
                    st.markdown("### 📊 Content Analysis")
                    
                    # Simple sentiment analysis for tone
                    try:
                        sentiment = vader.SentimentIntensityAnalyzer()
                        sent_scores = sentiment.polarity_scores(text)
                        compound = sent_scores.get('compound', 0.0)
                        if compound >= 0.3:
                            tone = 'Confident'
                        elif compound <= -0.2:
                            tone = 'Passive/Negative'
                        else:
                            tone = 'Neutral'
                    except:
                        tone = 'Neutral'
                        compound = 0.0
                    
                    st.markdown(f"**Sentiment Score:** {compound:.2f}")
                    st.markdown(f"**Tone Category:** {tone}")
                    
                    # Tone Category Metric
                    st.metric("Tone Category", tone, f"{compound:.2f}")

                    # Quick AI questions (use render_ai_message)
                    st.divider()
                    # Removed per request: AI Assistant for Skills section


                # ----- TAB 4: ATS Analysis -----
                with tab4:
                    st.markdown('<div class="section-header-modern fade-in-up">🎯 ATS Compatibility Analysis</div>', unsafe_allow_html=True)
                    col1, col2 = st.columns([1, 2])
                    with col1:
                        ats_score = results['ats_score']
                        st.metric("Overall ATS Score", f"{ats_score}/100")
                        st.progress(ats_score / 100)
                        if ats_score >= 80:
                            st.success("🎉 Excellent ATS compatibility")
                        elif ats_score >= 60:
                            st.warning("⚠️ Good ATS compatibility")
                        else:
                            st.error("🚨 Needs ATS optimization")
                    with col2:
                        st.markdown("### 💡 ATS Optimization Tips")
                        tips = [
                            "✅ Use standard section headings: Experience, Education, Skills, etc.",
                            "✅ Include relevant keywords naturally throughout your resume",
                            "✅ Use bullet points to improve readability and scanning",
                            "✅ Avoid images, graphics, tables, and complex formatting",
                            "✅ Use standard fonts like Arial, Calibri, or Times New Roman",
                            "✅ Save your resume as a PDF to preserve formatting",
                            "✅ Include your contact information prominently at the top",
                            "✅ Use consistent formatting throughout the document"
                        ]
                        for tip in tips:
                            st.markdown(tip)
                        # Removed per request: Missing keywords inside ATS tab


                # ----- TAB 5: AI Assistant (ChatGPT-like) -----
                with tab5:
                    st.markdown('<div class="section-header-modern fade-in-up">🤖 AI Resume Assistant</div>', unsafe_allow_html=True)
                    
                    # Check if RAG is available
                    if not analyzer.rag_enabled:
                        st.warning("⚠️ LangChain RAG AI Assistant requires OpenAI API key with credits. Please set OPENAI_API_KEY in environment variables and ensure you have credits.")
                    
                    # Initialize chat history if not exists
                    if 'ai_chat_history' not in st.session_state:
                        st.session_state.ai_chat_history = []
                    
                    # Display chat history
                    if st.session_state.ai_chat_history:
                        st.markdown("### 💬 Conversation History")
                        for message in st.session_state.ai_chat_history[-10:]:  # Show last 10 messages
                            if message["role"] == "user":
                                st.markdown(f"""
                                <div class="chat-message chat-user fade-in-up">
                                    <strong>👤 You:</strong>
                                    <div style="margin-top:0.5rem;">{message["content"]}</div>
                                </div>
                                """, unsafe_allow_html=True)
                            else:
                                render_ai_message(message["content"])
                    
                    # Chat input
                    st.markdown("### 🚀 Ask Questions About This Resume")
                    st.markdown("🔗 **Powered by:** LangChain + FAISS + OpenAI GPT-4o-mini")
                    st.markdown("💡 **Examples:** What certificates are mentioned? What projects are listed? What skills does the candidate have? What is the work experience?")
                    
                    # Quick action buttons for resume-specific questions
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button("📜 List Certificates", key="btn_certificates"):
                            prompt = "What certifications or certificates are mentioned in this resume? List them all with any details provided."
                            if analyzer.rag_enabled:
                                response = analyzer.rag_client.query(prompt, k=5)
                            else:
                                response = "AI Assistant is not available. Please check your OpenAI API key."
                            st.session_state.ai_chat_history.append({"role": "user", "content": prompt})
                            st.session_state.ai_chat_history.append({"role": "assistant", "content": response})
                            st.rerun()
                    
                    with col2:
                        if st.button("🚀 List Projects", key="btn_projects"):
                            prompt = "What projects are mentioned in this resume? Please list all projects with their descriptions and technologies used."
                            if analyzer.rag_enabled:
                                response = analyzer.rag_client.query(prompt, k=5)
                            else:
                                response = "AI Assistant is not available. Please check your OpenAI API key."
                            st.session_state.ai_chat_history.append({"role": "user", "content": prompt})
                            st.session_state.ai_chat_history.append({"role": "assistant", "content": response})
                            st.rerun()
                    
                    with col3:
                        if st.button("💼 Work Experience", key="btn_experience"):
                            prompt = "What is the candidate's work experience? List all positions, companies, durations, and key responsibilities mentioned."
                            if analyzer.rag_enabled:
                                response = analyzer.rag_client.query(prompt, k=5)
                            else:
                                response = "AI Assistant is not available. Please check your OpenAI API key."
                            st.session_state.ai_chat_history.append({"role": "user", "content": prompt})
                            st.session_state.ai_chat_history.append({"role": "assistant", "content": response})
                            st.rerun()
                    
                    # Custom question input
                    user_question = st.text_area(
                        "Ask me anything about this resume:",
                        placeholder="e.g., What programming languages does this candidate know? What is their educational background?",
                        height=100,
                        key="ai_chat_input"
                    )
                    
                    col1, col2 = st.columns([1, 4])
                    with col1:
                        send_clicked = st.button("Send 💬", type="primary", use_container_width=True, key="send_ai_chat")
                    with col2:
                        if st.button("Clear Chat 🗑️", use_container_width=True, key="clear_ai_chat"):
                            st.session_state.ai_chat_history = []
                            st.rerun()
                    
                    if send_clicked and user_question.strip():
                        # Use RAG to answer questions about the resume
                        if analyzer.rag_enabled:
                            with st.spinner("🤔 Analyzing resume content..."):
                                response = analyzer.rag_client.query(user_question, k=5)
                        else:
                            response = "🤖 LangChain RAG AI Assistant is currently unavailable. Please ensure:\n1. OPENAI_API_KEY is set in your environment with credits\n2. LangChain is installed (pip install langchain langchain-openai langchain-community)\n3. FAISS library is installed (pip install faiss-cpu)"
                        
                        st.session_state.ai_chat_history.append({"role": "user", "content": user_question})
                        st.session_state.ai_chat_history.append({"role": "assistant", "content": response})
                        st.rerun()


            except Exception as e:
                st.error(f"❌ Error during analysis: {str(e)}")
                st.error("Please check your resume format and try again.")
    else:
        st.markdown('<div class="section-header-modern fade-in-up">🚀 Getting Started</div>', unsafe_allow_html=True)
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            <div class="modern-card fade-in-up">
                <h3 style="color: #334155; margin-bottom: 1rem;">✨ Analysis Features</h3>
                <ul style="color: #64748b; line-height: 1.8;">
                    <li>🎯 ATS Compatibility Scoring</li>
                    <li>🔍 Skills Detection & Matching</li>
                    <li>🔑 Keyword Optimization</li>
                    <li>📊 Professional Visualizations</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div class="modern-card fade-in-up">
                <h3 style="color: #334155; margin-bottom: 1rem;">📄 Supported Formats</h3>
                <ul style="color: #64748b; line-height: 1.8;">
                    <li>📄 PDF Documents</li>
                    <li>📄 DOCX Files</li>
                    <li>📄 TXT Files</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)


# ---------------------------
# Main
# ---------------------------
def main():
    db = DatabaseManager()
    db.create_tables()


    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'user' not in st.session_state:
        st.session_state.user = None


    if not st.session_state.logged_in:
        login_page()
    else:
        dashboard_page()


if __name__ == "__main__":
    main()



